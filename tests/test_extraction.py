from pathlib import Path

from docx import Document
from pypdf import PdfWriter

from backend.processing.extractor import extract_text


def test_extract_text_from_txt(tmp_path: Path) -> None:
    file_path = tmp_path / "notes.txt"
    file_path.write_text("Hello from txt", encoding="utf-8")

    result = extract_text(file_path)

    assert result["success"] is True
    assert result["text"] == "Hello from txt"
    assert result["pages"] == 1
    assert result["characters"] == len("Hello from txt")


def test_extract_text_from_docx(tmp_path: Path) -> None:
    file_path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("First paragraph")
    document.add_paragraph("Second paragraph")
    document.save(file_path)

    result = extract_text(file_path)

    assert result["success"] is True
    assert result["text"] == "First paragraph\nSecond paragraph"
    assert result["pages"] == 1
    assert result["characters"] == len("First paragraph\nSecond paragraph")


def test_extract_text_from_pdf(tmp_path: Path) -> None:
    file_path = tmp_path / "sample.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with file_path.open("wb") as handle:
        writer.write(handle)

    result = extract_text(file_path)

    assert result["success"] is True
    assert result["pages"] == 1
    assert result["characters"] >= 0
    assert isinstance(result["text"], str)
